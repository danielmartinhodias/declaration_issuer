import streamlit as st
from datetime import datetime

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import os

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


from docx2pdf import convert

from datetime import datetime

# Get current timestamp
timestamp = datetime.now()


st.subheader ("Pedido de declaração de doença")

# st.button("**Declaração de doença**", use_container_width=True)
with st.form("my_form"):
    left, right=st.columns(2)

    nome = left.text_input("**Nome completo**")
    cc = right.text_input("Número de identificação (CC ou BI)")

    left, right=st.columns(2)

    inicio_ausencia=left.date_input("Primeiro dia de ausência ao trabalho", value=None)
    fim_ausencia=right.date_input("Último dia de ausência prevista ao trabalho", value=None)

    agree = st.checkbox("Confirmo, sob compromisso de honra, ter já previamente informado o médico do motivo de doença e ter sido avaliado pelo mesmo")
    submitted = st.form_submit_button("Submeter pedido")

# After form submission
if submitted:
  try:
    success_string=f"Eu, Daniel Martinho Ferreira Dias, portador da cédula profissional 63783, venho por este meio declarar que {nome}, portador do CC/BI {cc}, encontra-se doente, com incapacidade laboral desde {inicio_ausencia} e previsível até {fim_ausencia}. Por ser verdade e me ter sido pedido, emito a presente declaração que dato e assino"
    st.success(success_string)
    st.success("Enviado e-mail para o médico avaliar e validar.")

    doc = Document("./templates/declaracao_ausencia.docx")
    paragraph=doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run("Declaração de Doença")
    run.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run.bold = True                    # Make text bold
    run.font.name = 'Arial'           # Set font name
    run.font.size = Pt(16)            # Set font size (e.g., 14 pt)
    
    paragraph2=doc.add_paragraph(success_string)
    paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(f"./issued/dec_doenca{nome}.docx")

    convert(f"./issued/dec_doenca{nome}.docx")

      # Email credentials
    sender_email = "danielmartinhodias@gmail.com"
    receiver_email = "danielmartinhodias@gmail.com"
    app_password = "dsxxuvwybaevopdu"

    # Email content
    subject = f"Dec_doença_{nome}"
    body = f"Para editar, validar e assinar: {success_string}"

    # Create the email
    msg = MIMEMultipart()
    msg["From"] = sender_email
    msg["To"] = receiver_email
    msg["Subject"] = subject

    mensagem = f'Para validação de pedido de declaração de ausência de {nome}, submetida em {timestamp}'
    msg.attach(MIMEText(mensagem, 'plain'))

    anexos = [f"./issued/dec_doenca{nome}.docx", f"./issued/dec_doenca{nome}.pdf"]

    # Anexar arquivos
    for ficheiro in anexos:
        with open(ficheiro, 'rb') as f:
            parte = MIMEBase('application', 'octet-stream')
            parte.set_payload(f.read())
            encoders.encode_base64(parte)
            parte.add_header('Content-Disposition', f'attachment; filename="{os.path.basename(ficheiro)}"')
            msg.attach(parte)

    # Enviar e-mail

    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
        smtp.login(sender_email, app_password)
        smtp.send_message(msg)

    print("Email enviado com sucesso para o médico para validação. Aguarde contacto e resposta.")

  except:
    pass