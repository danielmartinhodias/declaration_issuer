from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx2pdf import convert

import forms_retrieval
import send_email

new_dict=forms_retrieval.get_answers("1wAjBWyt5yNi7gnVj4UD926svbfAIpAgmcecPzY_KJTg")

#definições

num=1

def get_variables(num):
  nome=list(list(new_dict["responses"][num]["answers"].values())[0]["textAnswers"]["answers"][0].values())[0]
  mcdts=list(list(new_dict["responses"][num]["answers"].values())[1]["textAnswers"]["answers"][0].values())[0]
  motivo=list(list(new_dict["responses"][num]["answers"].values())[2]["textAnswers"]["answers"][0].values())[0]
  subsistema=list(list(new_dict["responses"][num]["answers"].values())[3]["textAnswers"]["answers"][0].values())[0]
  numero_subsistema=list(list(new_dict["responses"][num]["answers"].values())[4]["textAnswers"]["answers"][0].values())[0]
  results=[nome,mcdts,motivo,subsistema,numero_subsistema]
  return (results)


nome=get_variables(num)[0]
subsistema=get_variables(num)[3]
numero_subsistema=get_variables(num)[4]
mcdts=get_variables(num)[1]


#escrita do documento

document = Document(('/Users/danielmdias/docs/_fun_time/00_madrid_prescriber/exams/code/data/template/template.docx'))

#styles
style1 = document.styles['Normal']
font = style1.font
font.name = 'Calibri'
font.size = Pt(16)


r  = "Eu, Daniel Martinho Dias, médico, portador da Cédula Profissional nº 63783, venho por este meio solicitar a {0}, {1} {2}, o(s) seguinte(s) MCDTs:".format(nome, subsistema, numero_subsistema)

p = document.add_paragraph(r)
p.paragraph_format.space_before = Pt(15)
p.paragraph_format.space_after = Pt(15)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.style = style1

print (nome)
# for i in mcdts:
p2=document.add_paragraph(mcdts)
p2
p2.style=style1


path='/Users/danielmdias/docs/_fun_time/00_madrid_prescriber/exams/code/data/output/Output_Test.docx'
document.save(path)


# convert(path)

# send_email.gmail_send_message_with_attachment("assina a requisição", '/Users/danielmdias/docs/_fun_time/00_madrid_prescriber/exams/code/data/output/Output_Test.docx')